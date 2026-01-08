"""
Transcript Formatter Module
===========================
Generate professionally formatted transcripts in TXT, DOCX, and PDF formats.
"""

import io
import re
import logging
from datetime import datetime
from typing import Optional

logger = logging.getLogger('transcriptai.transcript_formatter')


def _generate_title_from_filename(filename: Optional[str]) -> str:
    """Generate a clean title from filename or use default."""
    if not filename:
        return f"Transcript - {datetime.now().strftime('%B %d, %Y')}"

    # Remove extension and clean up
    name = re.sub(r'\.[^.]+$', '', filename)
    # Replace underscores and hyphens with spaces
    name = re.sub(r'[_-]+', ' ', name)
    # Title case
    name = name.strip().title()

    return name if name else f"Transcript - {datetime.now().strftime('%B %d, %Y')}"


def _format_text_content(text: str) -> str:
    """
    Format raw transcript text for professional presentation.
    Handles both speaker-labeled and plain text transcripts.
    """
    if not text or not text.strip():
        return ""

    lines = text.strip().split('\n')
    formatted_lines = []

    for line in lines:
        trimmed = line.strip()
        if not trimmed:
            formatted_lines.append("")
            continue

        # Check for speaker labels (e.g., "Speaker 1:", "[Speaker]:", ">>Speaker")
        speaker_match = re.match(r'^(>>|(?:\[?Speaker\s*\d*\]?:?))\s*(.*)$', trimmed, re.IGNORECASE)
        if speaker_match:
            speaker_text = speaker_match.group(2).strip()
            formatted_lines.append(f">> {speaker_text}" if speaker_text else "")
        else:
            formatted_lines.append(trimmed)

    return '\n'.join(formatted_lines)


def create_txt(text: str, title: Optional[str] = None, filename: Optional[str] = None) -> bytes:
    """
    Create a formatted plain text transcript.

    Args:
        text: Raw transcript text
        title: Optional title (auto-generated from filename if not provided)
        filename: Original filename for title generation

    Returns:
        UTF-8 encoded bytes of the formatted text
    """
    doc_title = title or _generate_title_from_filename(filename)
    formatted = _format_text_content(text)

    width = 70
    divider = "─" * width

    output = []
    output.append("")
    output.append("TRANSCRIPT".center(width))
    output.append(doc_title.center(width))
    output.append(datetime.now().strftime("%B %d, %Y").center(width))
    output.append("")
    output.append(divider)
    output.append("")

    for line in formatted.split('\n'):
        trimmed = line.strip()
        if not trimmed:
            output.append("")
        elif trimmed.startswith('>>'):
            output.append(f"  {trimmed[2:].strip()}")
        else:
            output.append(f"    {trimmed}")

    output.append("")
    output.append(divider)
    output.append("END OF TRANSCRIPT".center(width))
    output.append("")

    return '\n'.join(output).encode('utf-8')


def create_docx(text: str, title: Optional[str] = None, filename: Optional[str] = None) -> bytes:
    """
    Create a professionally formatted DOCX document.

    Args:
        text: Raw transcript text
        title: Optional title (auto-generated from filename if not provided)
        filename: Original filename for title generation

    Returns:
        DOCX file as bytes
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    doc_title = title or _generate_title_from_filename(filename)
    formatted = _format_text_content(text)
    subtitle = datetime.now().strftime("%B %d, %Y")

    # Colors
    NAVY = RGBColor(0x1A, 0x36, 0x5D)
    BLUE = RGBColor(0x31, 0x82, 0xCE)
    GRAY = RGBColor(0x2D, 0x37, 0x48)
    MUTED = RGBColor(0x71, 0x80, 0x96)
    DIVIDER = RGBColor(0xCB, 0xD5, 0xE0)

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header - Label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRANSCRIPT")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    # Header - Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(doc_title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Header - Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = GRAY

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("─" * 50)
    run.font.color.rgb = DIVIDER

    # Content
    for line in formatted.split('\n'):
        trimmed = line.strip()

        if not trimmed:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)

        if trimmed.startswith('>>'):
            p.paragraph_format.space_before = Pt(10)
            arrow = p.add_run("  ")
            arrow.font.color.rgb = BLUE
            arrow.font.size = Pt(11)
            text_run = p.add_run(trimmed[2:].strip())
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = GRAY
        else:
            p.paragraph_format.left_indent = Inches(0.25)
            text_run = p.add_run(trimmed)
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = GRAY

    # Footer divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("─" * 50)
    run.font.color.rgb = DIVIDER

    # Footer text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("END OF TRANSCRIPT")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info(f"Generated DOCX: {doc_title}")
    return buffer.getvalue()


def create_pdf(text: str, title: Optional[str] = None, filename: Optional[str] = None) -> bytes:
    """
    Create a professionally formatted PDF document.

    Args:
        text: Raw transcript text
        title: Optional title (auto-generated from filename if not provided)
        filename: Original filename for title generation

    Returns:
        PDF file as bytes
    """
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
    except ImportError:
        logger.error("fpdf2 not installed")
        raise ImportError("fpdf2 is required for PDF export. Install with: pip install fpdf2")

    doc_title = title or _generate_title_from_filename(filename)
    formatted = _format_text_content(text)
    subtitle = datetime.now().strftime("%B %d, %Y")

    # Colors
    NAVY = (26, 54, 93)
    BLUE = (49, 130, 206)
    GRAY = (45, 55, 72)
    MUTED = (113, 128, 150)
    DIVIDER = (203, 213, 224)

    class TranscriptPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.font_family = 'Helvetica'
            self.set_auto_page_break(auto=True, margin=25)
            self.set_margins(25, 25, 25)

        def footer(self):
            self.set_y(-15)
            self.set_font(self.font_family, '', 8)
            self.set_text_color(*MUTED)
            self.cell(0, 10, f'Page {self.page_no()}', align='C')

    pdf = TranscriptPDF()
    pdf.add_page()

    # Header - Label
    pdf.set_font(pdf.font_family, '', 10)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 10, 'TRANSCRIPT', align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Header - Title
    pdf.set_font(pdf.font_family, 'B', 20)
    pdf.set_text_color(*NAVY)
    pdf.cell(0, 12, doc_title, align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Header - Date
    pdf.set_font(pdf.font_family, 'I', 13)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 10, subtitle, align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(5)

    # Divider
    pdf.set_draw_color(*DIVIDER)
    pdf.set_line_width(0.5)
    pdf.line(40, pdf.get_y(), pdf.w - 40, pdf.get_y())
    pdf.ln(10)

    # Content
    for line in formatted.split('\n'):
        trimmed = line.strip()

        if not trimmed:
            pdf.ln(2)
            continue

        if trimmed.startswith('>>'):
            pdf.ln(3)
            pdf.set_font(pdf.font_family, '', 11)
            pdf.set_text_color(*GRAY)
            pdf.set_x(25)
            pdf.multi_cell(0, 6, trimmed[2:].strip())
        else:
            pdf.set_font(pdf.font_family, '', 11)
            pdf.set_text_color(*GRAY)
            pdf.set_x(32)
            pdf.multi_cell(0, 6, trimmed)

    # Footer divider
    pdf.ln(10)
    pdf.set_draw_color(*DIVIDER)
    pdf.line(40, pdf.get_y(), pdf.w - 40, pdf.get_y())
    pdf.ln(8)

    # Footer text
    pdf.set_font(pdf.font_family, '', 9)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 8, 'END OF TRANSCRIPT', align='C')

    # Output to bytes
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)

    logger.info(f"Generated PDF: {doc_title}")
    return buffer.getvalue()


def export_transcript(
    text: str,
    format: str = 'txt',
    title: Optional[str] = None,
    filename: Optional[str] = None
) -> tuple[bytes, str, str]:
    """
    Export transcript in the specified format.

    Args:
        text: Raw transcript text
        format: Output format ('txt', 'docx', 'pdf')
        title: Optional title
        filename: Original filename for title generation and output naming

    Returns:
        Tuple of (file_bytes, content_type, suggested_filename)
    """
    format = format.lower().strip()
    base_name = re.sub(r'\.[^.]+$', '', filename or 'transcript')

    if format == 'txt':
        content = create_txt(text, title, filename)
        return content, 'text/plain; charset=utf-8', f'{base_name}.txt'

    elif format == 'docx':
        content = create_docx(text, title, filename)
        return content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f'{base_name}.docx'

    elif format == 'pdf':
        content = create_pdf(text, title, filename)
        return content, 'application/pdf', f'{base_name}.pdf'

    else:
        raise ValueError(f"Unsupported format: {format}. Use 'txt', 'docx', or 'pdf'.")
